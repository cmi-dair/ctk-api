# pylint: disable=protected-access
import tempfile

import docx

from ctk_api.routers.summarization import anonymizer


def test_get_patient_name_success(document: tempfile._TemporaryFileWrapper) -> None:
    """Tests the get_patient_name function."""
    document = anonymizer.docx.Document(document)
    first_name, last_name = anonymizer.get_patient_name(document)

    assert first_name == "Lea"
    assert last_name == "Avatar"


def test_get_diagnostic_paragraphs(document: tempfile._TemporaryFileWrapper) -> None:
    """Tests the get_diagnostic_paragraphs function."""
    document = anonymizer.docx.Document(document)

    diagnostic_paragraphs = anonymizer.get_diagnostic_paragraphs(document)

    assert len(diagnostic_paragraphs) == 3
    assert diagnostic_paragraphs[0].text == "clinical summary and impression"


def test_find_and_replace_no_match_case() -> None:
    """Tests the find_and_replace function without matching case."""
    text = "This is Is a test sentence."
    target = "is"
    replacement = "rookie harbor"
    expected = "This rookie harbor rookie harbor a test sentence."

    actual = anonymizer._find_and_replace(text, target, replacement)

    assert actual == expected


def test_find_and_replace_match_case() -> None:
    """Tests the find_and_replace function with matching case."""
    text = "This is Is a test sentence."
    target = "is"
    replacement = "rookie harbor"
    expected = "This rookie harbor Rookie harbor a test sentence."

    actual = anonymizer._find_and_replace(text, target, replacement, match_case=True)

    assert actual == expected


def test_find_and_replace_match_case_and_slash() -> None:
    """Tests the find_and_replace function with matching case and slash."""
    text = "This is Is a test sentence."
    target = "is"
    replacement = "rookie/harbor"
    expected = "This rookie/harbor Rookie/Harbor a test sentence."

    actual = anonymizer._find_and_replace(text, target, replacement, match_case=True)

    assert actual == expected


def test_anonymize_paragraphs(document: tempfile._TemporaryFileWrapper) -> None:
    """Tests whether paragraphs are anonymized correctly."""
    doc = docx.Document(document.name)
    first_name, last_name = anonymizer.get_patient_name(doc)

    actual = anonymizer.anonymize_paragraphs(doc.paragraphs, first_name, last_name)

    assert len(actual) == len(doc.paragraphs)
    assert actual[0].text == "Title"
    assert actual[1].text == "clinical summary and impression"
    assert actual[2].text == "Name: [FIRST_NAME] [LAST_NAME]"
    assert actual[3].text == "He/She he/she himself/herself man/woman"
