"""Configurations for all tests."""
import tempfile
from collections.abc import Generator

import docx
import pytest

from ctk_api.core import config
from ctk_api.microservices import elastic

settings = config.get_settings()
ELASTIC_DIAGNOSES_INDEX = settings.ELASTIC_DIAGNOSES_INDEX
ELASTIC_SUMMARIZATION_INDEX = settings.ELASTIC_SUMMARIZATION_INDEX


@pytest.fixture()
def document() -> Generator[tempfile._TemporaryFileWrapper, None, None]:
    """Generates a fixture Word document.

    This document contains a title and a paragraph.

    Returns:
        Generator[tempfile._TemporaryFileWrapper, None, None]: The saved file.
    """
    doc = docx.Document()
    doc.add_heading("Title", 0)
    doc.add_heading("clinical summary and impressions", 1)
    doc.add_paragraph("Name: Lea Avatar")
    doc.add_paragraph("He she herself man")

    with tempfile.NamedTemporaryFile(suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        yield temp_file


@pytest.fixture(scope="session")
def elastic_client() -> elastic.ElasticClient:
    """Returns an Elasticsearch client."""
    return elastic.ElasticClient()


@pytest.fixture(scope="session", autouse=True)
def _clear_elastic(
    elastic_client: elastic.ElasticClient,
) -> None:
    """Clears the test Elasticsearch indices."""
    if not ELASTIC_DIAGNOSES_INDEX.startswith(
        "test_",
    ) or not ELASTIC_SUMMARIZATION_INDEX.startswith("test_"):
        msg = "Testing Elasticsearch indices must start with 'test_'."
        raise ValueError(msg)
    elastic_client.client = elastic_client.client.options(ignore_status=[400, 404])
    elastic_client.client.indices.delete(index=ELASTIC_DIAGNOSES_INDEX)
    elastic_client.client.indices.delete(index=ELASTIC_SUMMARIZATION_INDEX)
