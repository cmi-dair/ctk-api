# pylint: disable=redefined-outer-name
import tempfile
from typing import Generator

import docx
import pytest
from fastapi import status, testclient

from . import conftest


@pytest.fixture
def document() -> Generator[tempfile._TemporaryFileWrapper, None, None]:
    """Generates a new Word document with a title and a paragraph, and returns
    the path to the saved file.

    Returns:
        tempfile.NamedTemporaryFile: The saved file.
    """
    doc = docx.Document()
    doc.add_heading("Title", 0)
    doc.add_heading("clinical summary and impression", 1)
    doc.add_paragraph("Name: Lea Avatar")
    doc.add_paragraph("He she herself man")

    with tempfile.NamedTemporaryFile(suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        yield temp_file


def test_anonymization_endpoint(
    client: testclient.TestClient,
    endpoints: conftest.Endpoints,
    document: tempfile._TemporaryFileWrapper,
) -> None:
    """Tests the anonymization endpoint."""
    form_data = {"docx_file": document}
    expected = "\n".join(
        [
            "clinical summary and impression",
            "Name: [FIRST_NAME] [LAST_NAME]",
            "He/She he/she himself/herself man/woman",
        ]
    )

    response = client.post(endpoints.ANONYMIZE_REPORT, files=form_data)

    assert response.status_code == status.HTTP_200_OK
    assert response.json() == expected
